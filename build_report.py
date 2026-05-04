"""
Generates the comprehensive Word document for the India NFHS-5 ECD project.
Run from /Users/prachi/india_nfhs5_project/
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIG = "/Users/prachi/india_nfhs5_project/figures"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin = section.right_margin = Cm(2.54)
section.top_margin  = section.bottom_margin = Cm(2.54)

# ── Styles helper functions ───────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    return p

def add_figure(doc, path, caption, width_in=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Figure not found: {path}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "BDD7EE")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx+1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table

def formula_para(doc, formula_text):
    """Add a centred formula paragraph."""
    p = doc.add_paragraph(formula_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(11)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading("Modelling Early Childhood Development\nusing the UNICEF ECDI2030 Framework", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("India Analysis — NFHS-5 (2019–21) Proxy Pipeline")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].bold = True

doc.add_paragraph("")
meta = doc.add_paragraph(
    "Author: Prachi\nMaster's Thesis Project (MTP2)\n"
    "Data Source: DHS-7 / NFHS-5 Kids Recode (IAKR7EFL.DTA)\n"
    "Analysis Date: April 2026\n"
    "GitHub: github.com/ps12-1/Prachi_MTP2_India_NFHS5"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in meta.runs:
    run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "Table of Contents")
toc_items = [
    "1. Introduction and Motivation",
    "2. The UNICEF ECDI2030 Framework",
    "3. Data Source: India NFHS-5",
    "4. Methodological Challenge: No Direct ECDI Items in NFHS-5",
    "5. Proxy Indicator Construction (Expanded 4-Domain Composite)",
    "6. Descriptive Results — ECD Proxy Rates",
    "7. Supervised Machine Learning: Predicting ECD Proxy Composite",
    "8. Unsupervised Machine Learning: Clustering Developmental Profiles",
    "9. Geospatial Analysis: State-Level Patterns and Spatial Autocorrelation",
    "10. Econometric Analysis: Core Models (10.1–10.6) and Robustness Suite (10.7–10.12)",
    "11. Figures Explained",
    "12. Limitations",
    "13. Conclusions and Policy Implications",
    "References",
]
for item in toc_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: Introduction
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "1. Introduction and Motivation")

body(doc,
    "Early childhood development (ECD) — encompassing cognitive, physical, socio-emotional, and learning "
    "dimensions of a child's growth in the first five years of life — is widely recognised as one of the most "
    "critical determinants of long-run human capital formation. Investments made during this period have been "
    "shown to yield higher returns than at any later stage of the life course (Heckman, 2006). SDG Target 4.2.1 "
    "specifically mandates that all children have access to quality early childhood development, care, and "
    "pre-primary education by 2030.")

body(doc,
    "India presents a particularly important context for ECD research. With over 90 million children under five "
    "years of age, India accounts for roughly one-fifth of the global burden of child stunting and faces stark "
    "geographic and socioeconomic inequalities in child outcomes. The Integrated Child Development Services (ICDS) "
    "programme, delivered through a nationwide network of Anganwadi Centres (AWCs), constitutes the world's "
    "largest ECD programme by reach.")

body(doc,
    "This project applies machine learning and geospatial methods to the India National Family Health Survey "
    "(NFHS-5, 2019–21) to model ECD outcomes across India's 36 states and Union Territories. A key "
    "methodological contribution of this work is the construction of a proxy ECDI2030 composite from NFHS-5 "
    "variables, given that the survey does not administer direct cognitive or socio-emotional assessments for "
    "children.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: ECDI2030 Framework
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "2. The UNICEF ECDI2030 Framework")

body(doc,
    "The Early Childhood Development Index 2030 (ECDI2030) is an internationally standardised tool developed by "
    "UNICEF to measure developmentally appropriate milestones for children aged 24–59 months. It was released in "
    "September 2023 (UNICEF Technical Manual, 2023) and replaces the original ECDI used in MICS4/5 surveys.")

heading2(doc, "2.1 Domains and Items")

body(doc,
    "ECDI2030 measures 10 binary items across 4 developmental domains. Each item is coded 1 (pass) or 0 (fail). "
    "Items EC14 and EC15 are reverse-coded — a response of 'Yes' (1) indicates a negative behaviour and is "
    "therefore a FAIL.")

add_table(doc,
    ["Domain", "Item Code", "Description", "Pass Condition"],
    [
        ["Literacy-Numeracy", "EC6", "Can identify/name 10+ alphabet letters", "Response = 1 (Yes)"],
        ["Literacy-Numeracy", "EC7", "Can read 4+ simple, popular words", "Response = 1 (Yes)"],
        ["Literacy-Numeracy", "EC8", "Knows names & symbols for all numbers 1–10", "Response = 1 (Yes)"],
        ["Physical", "EC9", "Fine motor: picks up small object with two fingers", "Response = 1 (Yes)"],
        ["Physical", "EC10", "Energy/health: not too sick to play", "Response = 1 (Yes)"],
        ["Learning", "EC11", "Can follow two-step directions/instructions", "Response = 1 (Yes)"],
        ["Learning", "EC12", "Picks up and plays with things (attention, curiosity)", "Response = 1 (Yes)"],
        ["Socio-emotional", "EC13", "Gets along well with other children", "Response = 1 (Yes)"],
        ["Socio-emotional", "EC14 (R)", "Does NOT kick/bite/hit other children", "Response = 2 (No) — REVERSE"],
        ["Socio-emotional", "EC15 (R)", "Is NOT easily distracted", "Response = 2 (No) — REVERSE"],
    ]
)

doc.add_paragraph("")
heading2(doc, "2.2 Domain Scoring")
body(doc,
    "A child is considered 'on track' for a domain if and only if they pass ALL items within that domain:")
body(doc,
    "  Literacy-Numeracy: on track if EC6 = EC7 = EC8 = 1\n"
    "  Physical:          on track if EC9 = EC10 = 1\n"
    "  Learning:          on track if EC11 = EC12 = 1\n"
    "  Socio-emotional:   on track if EC13 = 1 AND EC14 = 2 AND EC15 = 2")

heading2(doc, "2.3 Composite Score")
body(doc,
    "The ECDI2030 composite ('developmentally on track') is defined as:")
formula_para(doc, "Composite = 1  if  domains_on_track ≥ 3  out of 4")
body(doc,
    "This is a conservative threshold: a child must be on track in at least three of the four developmental "
    "domains to be classified as developmentally on track overall. If any domain score is missing, the composite "
    "is treated as missing (listwise deletion, not imputed).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: Data Source
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "3. Data Source: India NFHS-5 (2019–21)")

body(doc,
    "The National Family Health Survey, Fifth Round (NFHS-5) was conducted in India from 2019 to 2021. It is "
    "India's implementation of the Demographic and Health Survey (DHS) programme (DHS-7), overseen by the "
    "International Institute for Population Sciences (IIPS), Mumbai, and funded by USAID.")

heading2(doc, "3.1 Survey Design")
body(doc,
    "NFHS-5 employed a stratified, two-stage cluster sampling design. In the first stage, primary sampling "
    "units (PSUs) — villages in rural areas and Census Enumeration Blocks (CEBs) in urban areas — were selected "
    "using probability proportional to size (PPS) sampling. In the second stage, households were selected using "
    "systematic random sampling within each PSU. The survey covered all 36 states and Union Territories.")

add_table(doc,
    ["Parameter", "Value"],
    [
        ["Survey round", "NFHS-5 / DHS-7"],
        ["Field period", "2019–2021"],
        ["Total KR file rows", "232,920 birth records"],
        ["Analysis sample (alive, 24–59 months)", "133,750 children"],
        ["States/UTs covered", "36"],
        ["Sampling design", "Stratified 2-stage cluster PPS"],
        ["Sample weight variable", "v005 (divide by 1,000,000)"],
        ["Primary data file used", "IAKR7EFL.DTA (Kids Recode, Stata format)"],
    ]
)

heading2(doc, "3.2 Kids Recode (KR) File Structure")
body(doc,
    "The Kids Recode file (IAKR7EFL.DTA) contains one row per child born in the five years preceding the survey. "
    "It incorporates both the birth history data (b* variables) and the child health module data (h*, hw*, s5* "
    "variables). Mother-level variables (v* variables) are merged into each birth record, so no separate merge "
    "with the Individual Recode (IR) file is required for our core analysis.")

heading2(doc, "3.3 Key Variables Used")
add_table(doc,
    ["Variable", "Description", "Use in Analysis"],
    [
        ["hw1", "Child age in months (measurement roster)", "Age filter and predictor"],
        ["b4", "Sex of child (male/female)", "Predictor"],
        ["b5", "Child alive? (yes/no)", "Inclusion filter"],
        ["hw70", "Height-for-age z-score (HAZ × 100)", "Physical proxy outcome"],
        ["hw71", "Weight-for-age z-score (WAZ × 100)", "Nutritional indicator"],
        ["hw72", "Weight-for-height z-score (WHZ × 100)", "Descriptive"],
        ["s558", "Received anganwadi/ICDS benefit last 12 months (yes/no)", "Selection into anganwadi module"],
        ["s562", "Frequency of anganwadi attendance (regularly/occasionally/not at all)", "Learning proxy outcome"],
        ["v024", "State (36 states/UTs)", "Geospatial analysis; state dummies"],
        ["v025", "Urban/rural residence", "Predictor"],
        ["v190", "Wealth index combined quintile (poorest–richest)", "Predictor"],
        ["v149", "Mother's educational attainment (no education–higher)", "Predictor"],
        ["v012", "Mother's current age in years", "Predictor"],
        ["v005", "Women's sample weight (÷ 1,000,000 for probability weight)", "Descriptive weighting note"],
    ]
)

heading2(doc, "3.4 Anthropometric Z-Score Encoding")
body(doc,
    "NFHS-5/DHS stores anthropometric z-scores as integers scaled by a factor of 100. The conversion is:")
formula_para(doc, "HAZ_actual = hw70 / 100")
formula_para(doc, "WAZ_actual = hw71 / 100")
body(doc,
    "For example, hw70 = −200 corresponds to HAZ = −2.0 SD; hw70 = 147 corresponds to HAZ = +1.47 SD. "
    "Values outside the plausible range [−600, +600] (i.e., z-scores outside ±6 SD) are flagged as implausible "
    "by DHS and treated as missing in this analysis. The plausible range check follows WHO recommendations "
    "(WHO, 2006).")

body(doc,
    "Z-scores are computed using the WHO 2006 Multicentre Growth Reference Standards, which establish sex- "
    "and age-specific reference populations derived from children raised under optimal conditions across six "
    "countries (Brazil, Ghana, India, Norway, Oman, United States).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: Methodological Challenge
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "4. Methodological Challenge: Absence of Direct ECDI Items in NFHS-5")

body(doc,
    "The central methodological challenge of this project is that NFHS-5 was not designed as a child "
    "development assessment survey. Unlike the UNICEF MICS6 instrument (which directly administers EC6–EC15 "
    "items to children 24–59 months), NFHS-5 focuses on maternal and child health, fertility, nutrition, and "
    "health service access. It contains no cognitive tests, motor skill assessments, or socio-emotional "
    "observations for children.")

body(doc,
    "To address this gap we adopt a small-area integration strategy: we merge external household-survey data "
    "into the NFHS-5 child record using the demographic key v024 (state of residence). Each child inherits "
    "the state-level value of an externally measured indicator as their estimated value for the unobserved "
    "ECD domain. This approach is standard in cross-survey integration when individual-level items are "
    "unobservable (Pfeffermann 2013; Rao & Molina 2015). Two external sources are used:")
body(doc,
    "1. ASER 2019 'Early Years' state-level proportions of children aged 4–5 enrolled in any pre-school + "
    "passing the picture-description (early language) and counting-objects (early numeracy) tasks. Used as "
    "a Literacy–Numeracy proxy (ECDI2030 EC6–EC8). Source: ASER Centre, India. The values are extracted "
    "directly from the public ASER 2019 Early Years PDF report (196 pages) using a pdfplumber-based parser "
    "(external_data/parse_aser_2019.py). ASER 2019 sampled 26 districts in 24 states (one rural district "
    "per state, with two districts in MP and UP). For the 12 states/UTs not sampled by ASER (Andaman & "
    "Nicobar, Arunachal Pradesh, Chandigarh, Dadra & Nagar Haveli & Daman & Diu, Goa, Jammu & Kashmir, "
    "Ladakh, Lakshadweep, Mizoram, NCT of Delhi, Puducherry, Sikkim) the value is imputed by the regional "
    "ASER mean (north / south / east / west / central / north-east).")
body(doc,
    "2. NFHS-5 published State Compendium of Factsheets (IIPS 2021) — state-level shares of married women "
    "in 3+ household decisions (#119), women employed in last 12 months (#120), and ever-married women "
    "experiencing spousal violence (#125, reverse-coded). Combined into a composite 'socio-emotional "
    "household environment' index used as a proxy for ECDI2030 EC13–EC15. Source: IIPS NFHS-5 State "
    "Factsheet Compendium PDFs (Phase I + Phase II). The CSV is sourced from the community repository "
    "github.com/jvargh7/nfhs5_factsheets, which extracts the indicator tables from the Phase-I and Phase-II "
    "compendium PDFs into structured form. Coverage: all 36 states/UTs.")
body(doc,
    "The following table summarises the proxy availability for each ECDI2030 item under this expanded "
    "framework:")

add_table(doc,
    ["ECDI2030 Domain", "Item", "What It Measures", "Proxy Source", "Coverage"],
    [
        ["Literacy-Numeracy", "EC6", "Identifies 10+ alphabet letters",
         "ASER 2019 letter-recognition rate (state)", "100% (state-merged)"],
        ["Literacy-Numeracy", "EC7", "Reads 4+ simple words",
         "ASER 2019 letter+number aggregate (state)", "100% (state-merged)"],
        ["Literacy-Numeracy", "EC8", "Knows numbers 1–10",
         "ASER 2019 number-recognition rate (state)", "100% (state-merged)"],
        ["Physical", "EC9", "Fine motor", "None available — note", "—"],
        ["Physical", "EC10", "Not too sick to play",
         "HAZ > −2 SD (hw70) — PARTIAL", "94%"],
        ["Learning", "EC11", "Follows 2-step directions", "None available — note", "—"],
        ["Learning", "EC12", "Picks up/plays with things",
         "Regular anganwadi attendance (s562)", "69%"],
        ["Socio-emotional", "EC13", "Gets along with other children",
         "NFHS-5 women's empowerment composite (state)", "100% (state-merged)"],
        ["Socio-emotional", "EC14 (R)", "Does not kick/bite/hit",
         "NFHS-5 spousal-violence indicator, reverse (state)", "100% (state-merged)"],
        ["Socio-emotional", "EC15 (R)", "Not easily distracted",
         "NFHS-5 women's decision-making (state)", "100% (state-merged)"],
    ]
)

body(doc,
    "Because the Literacy-Numeracy and Socio-emotional proxies are merged at the state level, every child "
    "inherits a value (no missingness from the external sources). However, this introduces an ecological-"
    "fallacy risk: children within a state share the same proxy value, so within-state variation in the "
    "missing domain is by construction zero. To partially mitigate this, the proxy domain definitions in "
    "Section 5 require BOTH a state-level signal AND an individual-level corroborator (mother's education "
    "for Literacy-Numeracy; non-poorest wealth for Socio-emotional).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: Proxy Construction
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "5. Proxy Indicator Construction (Expanded 4-Domain Composite)")

heading2(doc, "5.1 Physical Domain Proxy: Not Stunted")

body(doc,
    "Stunting — defined as height-for-age z-score (HAZ) below −2 standard deviations from the WHO 2006 "
    "reference median — is the most widely used indicator of chronic undernutrition and has been extensively "
    "validated as a proxy for physical developmental status in children under five (Black et al., 2013; "
    "Victora et al., 2008). A child who is not stunted (HAZ > −2 SD) is classified as 'on track' for the "
    "physical proxy domain.")

formula_para(doc, "ecdi_proxy_physical = 1  if  HAZ > −2.0 SD,  else 0")
formula_para(doc, "                    = NaN  if  HAZ measurement missing")

body(doc,
    "Note: This proxy captures nutritional history and physical health rather than the fine motor skill (EC9) "
    "and energy/health (EC10) items in the official ECDI2030 Physical domain. It is a structural proxy, not a "
    "direct measurement.")

heading2(doc, "5.2 Nutritional Domain (Additional Indicator, Not in Composite)")

body(doc,
    "Weight-for-age z-score (WAZ) < −2 SD defines underweight status. This is reported as a separate "
    "descriptive indicator but not included in the composite to avoid double-counting physical health:")
formula_para(doc, "ecdi_proxy_nutritional = 1  if  WAZ > −2.0 SD,  else 0")

heading2(doc, "5.3 Learning Domain Proxy: Regular Anganwadi Attendance")

body(doc,
    "India's ICDS programme delivers structured early learning activities through Anganwadi Centres (AWCs). "
    "NFHS-5 variable s562 records how frequently a child attended the AWC in the previous 12 months: "
    "'regularly', 'occasionally', 'not at all', or unknown. Regular attendance is used as a proxy for "
    "consistent exposure to a structured early learning environment, analogous to the learning domain's "
    "emphasis on sustained engagement and attention:")
formula_para(doc, "ecdi_proxy_learning = 1  if  s562 = 'regularly',  else 0")
formula_para(doc, "                    = NaN  if  s558 = 'no'  (child never enrolled)")

body(doc,
    "Important caveat: s562 is only populated for children whose mother reported that the child received "
    "some anganwadi benefit in the past 12 months (s558 = 'yes'). Children whose mothers report no benefit "
    "(s558 = 'no') are assigned NaN for s562 — they are not recorded as 'not attending' but rather as outside "
    "the anganwadi system. This creates a selection effect: the learning proxy only captures ECCE access among "
    "those already enrolled in the ICDS programme.")

heading2(doc, "5.4 Literacy–Numeracy Domain Proxy (Individual-Level Score)")

body(doc,
    "Earlier versions of this proxy used a binary state-merged rule, where every child in the same state "
    "inherited the same Literacy-Numeracy classification. That construction made within-state variation "
    "in the domain identically zero, exposing the proxy to ecological-fallacy critiques. We now use an "
    "INDIVIDUAL-LEVEL CONTINUOUS SCORE that combines the state-level ASER signal with within-state "
    "variation drawn from NFHS-5 KR covariates. Following Fay-Herriot (1979) small-area estimation, each "
    "child's latent literacy-numeracy score is:")
formula_para(doc, "litnum_score_ij = z(state_aser_rate_state(i))")
formula_para(doc, "                + 0.40 · z(mother_edu_ij)")
formula_para(doc, "                + 0.30 · z(wealth_index_ij)")
formula_para(doc, "                + 0.10 · is_urban_ij")
body(doc,
    "where z(·) denotes the standardised value of the variable. The β-weights are calibrated against the "
    "Black et al. (2017) Lancet ECD series, in which mother's education and household material resources "
    "are the two strongest individual-level correlates of cognitive-stimulation outcomes. The score is "
    "z-rescaled across the analytic sample, and the binary on-track classification is set at score ≥ 0 "
    "(above-median individual). Within-state SD of the resulting score is 0.421 (mean across 36 states) — "
    "non-zero by construction. A continuous CDF transform Φ(litnum_score) provides a 0–1 probability used "
    "in Section 10's quantile and IV specifications.")
formula_para(doc, "ecdi_proxy_litnum = 1  if  litnum_score ≥ 0,   else 0")

heading2(doc, "5.5 Socio-emotional Domain Proxy (Individual-Level Score)")
body(doc,
    "Symmetrically, the Socio-emotional proxy is built as an individual-level continuous score that "
    "combines the state-level women's-empowerment signal with within-state variation from caregiver "
    "wealth, education, and age:")
formula_para(doc, "se_score_ij = z(socioemo_state_index_state(i))")
formula_para(doc, "             + 0.35 · z(wealth_index_ij)")
formula_para(doc, "             + 0.30 · z(mother_edu_ij)")
formula_para(doc, "             + 0.15 · z(mother_age_ij)")
body(doc,
    "The β-weights reflect Walker et al. (2011) Lancet ECD findings on caregiver-responsiveness "
    "predictors. Within-state SD of the resulting score is 0.393 (mean across states). On-track threshold "
    "is again the median (score ≥ 0).")
formula_para(doc, "ecdi_proxy_se = 1  if  se_score ≥ 0,   else 0")
add_figure(doc, f"{FIG}/18_within_state_litnum.png",
    "Figure 18. Within-state distribution of the individual Literacy-Numeracy score, six largest states. "
    "Demonstrates non-zero within-state variation — ecological-fallacy concern is mitigated.")

heading2(doc, "5.6 Four-Domain ECD Proxy Composite")

body(doc,
    "Following the official UNICEF ECDI2030 scoring rule (a child is 'developmentally on track' if they pass "
    "≥ 3 of 4 domains), we construct the expanded composite as:")
formula_para(doc, "ecdi_proxy4_composite = 1  if  domains_on_track ≥ 3 / 4")
formula_para(doc, "                      = 0  otherwise")

body(doc,
    "where domains_on_track is the sum of ecdi_proxy_physical, ecdi_proxy_learning, ecdi_proxy_litnum, and "
    "ecdi_proxy_se. The composite is missing when any domain value is missing (in practice, only when HAZ or "
    "anganwadi data are missing, since the externally-merged Literacy-Numeracy and Socio-emotional values are "
    "complete for all 36 states).")
body(doc,
    "Compared to the 2-domain composite (25.0%), the 4-domain composite has a higher overall on-track rate "
    "(44.6% with real parsed ASER + NFHS-5 factsheet data) because a child can satisfy 3 of 4 domains, "
    "including the externally-merged environmental indicators, even if they fail one of the two NFHS-5-only "
    "domains. The 4-domain composite covers the same n = 86,999 children with complete physical + learning "
    "data.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: Descriptive Results
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "6. Descriptive Results — ECD Proxy Rates")

heading2(doc, "6.1 Sample Characteristics")

add_table(doc,
    ["Characteristic", "Value"],
    [
        ["Total children 24–59 months (alive)", "133,750"],
        ["Female (%)", "48.2%"],
        ["Urban (%)", "20.9%"],
        ["With HAZ measurement (%)", "94.4% (n = 126,250)"],
        ["With WAZ measurement (%)", "95.2% (n = 127,318)"],
        ["With anganwadi attendance data (%)", "68.6% (n = 91,719)"],
        ["With complete proxy composite (%)", "65.0% (n = 86,999)"],
        ["Mean HAZ (SD)", "−1.47 SD"],
        ["Mean WAZ (SD)", "−1.47 SD"],
        ["Wealth distribution", "Poorest 27% / Poorer 23% / Middle 19% / Richer 17% / Richest 13%"],
        ["Largest states by sample", "UP (20,034), Bihar (11,612), MP (9,086), Rajasthan (8,500)"],
    ]
)

heading2(doc, "6.2 Proxy Domain On-Track Rates")

body(doc, "Two-domain composite (NFHS-5 only — original specification):")
add_table(doc,
    ["Proxy Domain", "On-Track Rate", "Sample Size", "Interpretation"],
    [
        ["Physical (not stunted, HAZ > −2 SD)", "62.2%", "126,250", "~38% of 24–59 month children are stunted"],
        ["Nutritional (not underweight, WAZ > −2 SD)", "67.2%", "127,318", "~33% are underweight"],
        ["Learning (regular anganwadi attendance)", "40.7%*", "91,719", "* Among enrolled children only"],
        ["2-domain Composite (physical + learning)", "25.0%", "86,999", "Joint probability"],
    ]
)
body(doc, "")
body(doc, "Expanded 4-domain composite (individual-level scores, real merged data):")
add_table(doc,
    ["Proxy Domain", "On-Track Rate", "Sample Size", "Construction"],
    [
        ["Physical (not stunted)",                         "62.2%", "126,250", "NFHS-5 KR (HAZ)"],
        ["Learning (regular anganwadi)",                   "40.7%",  "91,719", "NFHS-5 KR (s562)"],
        ["Literacy-Numeracy (individual score ≥ 0)",       "50.1%",  "86,999", "ASER 2019 + mother edu + wealth + urban"],
        ["Socio-emotional (individual score ≥ 0)",         "46.8%",  "86,999", "NFHS-5 factsheet + wealth + edu + age"],
        ["4-domain Composite (≥ 3 of 4)",                  "35.8%",  "86,999", "ECDI2030 official scoring rule"],
    ]
)
body(doc, "")
body(doc,
    "Figure 1 (01_proxy_domain_rates.png) shows the original 4 proxy rates from the 2-domain analysis. "
    "Figure 14 (14_expanded_domain_rates.png) shows the 5 rates of the expanded 4-domain construction "
    "with individual-level scores, where the composite rate is 35.8% (survey-weighted 36.9%). Within-state "
    "SD of the individual Literacy-Numeracy score is 0.421 and of the Socio-emotional score is 0.393 — "
    "both strictly positive, confirming the ecological-fallacy mitigation in Section 5.")

heading2(doc, "6.3 Variation by Background Characteristics")

body(doc, "Proxy composite on-track rates by key determinants:")

add_table(doc,
    ["Subgroup", "On-Track Rate"],
    [
        ["Female children", "25.5%"],
        ["Male children", "24.5%"],
        ["Urban", "27.0%"],
        ["Rural", "24.5%"],
        ["Poorest quintile", "20.8%"],
        ["Poorer quintile", "24.3%"],
        ["Middle quintile", "27.4%"],
        ["Richer quintile", "28.9%"],
        ["Richest quintile", "26.7%"],
        ["No education (mother)", "20.6%"],
        ["Incomplete secondary (mother)", "26.5%"],
        ["Higher education (mother)", "28.4%"],
        ["Complete secondary (mother)", "30.6%"],
    ]
)

body(doc,
    "A consistent wealth gradient is observed from the poorest (20.8%) to the richer quintile (28.9%), "
    "though the richest quintile (26.7%) shows a slight dip compared to richer — possibly reflecting that "
    "wealthy urban families bypass anganwadi in favour of private nursery programmes. Mother's education "
    "shows a consistent protective gradient from no education (20.6%) to complete secondary (30.6%), but "
    "the higher education group (28.4%) is lower than complete secondary — again possibly reflecting "
    "private preschool substitution.")

heading2(doc, "6.4 State-Level Variation")

body(doc,
    "The ten highest and lowest states are presented below. "
    "Figure 5 (06_state_rates.png) shows the full ranked bar chart.")

add_table(doc,
    ["Rank", "State", "On-Track Rate"],
    [
        ["1 (Highest)", "Andaman & Nicobar Islands", "47.1%"],
        ["2", "West Bengal", "43.2%"],
        ["3", "Odisha", "41.7%"],
        ["4", "Dadra & Nagar Haveli and Daman & Diu", "41.3%"],
        ["5", "Andhra Pradesh", "40.0%"],
        ["6", "Gujarat", "39.8%"],
        ["...", "...", "..."],
        ["31", "Uttar Pradesh", "16.6%"],
        ["32", "Chandigarh", "15.5%"],
        ["33", "Assam", "15.1%"],
        ["34", "Arunachal Pradesh", "14.1%"],
        ["35", "Meghalaya", "8.5%"],
        ["36 (Lowest)", "Nagaland", "0.3%"],
    ]
)

body(doc,
    "The Northeast states of Nagaland, Manipur, and Meghalaya show very low composite rates, "
    "primarily driven by near-zero anganwadi attendance in the survey. These states likely rely on "
    "alternative ECD delivery systems (church-based crèches, community centres) that are not captured "
    "by NFHS-5's anganwadi-specific questions, making the learning proxy unreliable for these states. "
    "This represents a limitation of the proxy approach, not necessarily a true deficit in early "
    "childhood learning in these regions.")

add_figure(doc, f"{FIG}/01_proxy_domain_rates.png",
    "Figure 1. ECD Proxy Domain On-Track Rates — India NFHS-5 (24–59 months)")
add_figure(doc, f"{FIG}/02_anthropometric_distributions.png",
    "Figure 2. HAZ and WAZ Distributions — NFHS-5 India. Red dashed line = −2 SD stunting/underweight threshold.",
    width_in=6.0)
add_figure(doc, f"{FIG}/03_composite_by_wealth.png",
    "Figure 3. ECD Proxy Composite On-Track Rate by Wealth Quintile — India NFHS-5.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7: Supervised ML
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "7. Supervised Machine Learning: Predicting ECD Proxy Composite")

heading2(doc, "7.1 Prediction Task")

body(doc,
    "The supervised learning task is to predict whether a child is 'on track' on the proxy composite "
    "(ecdi_proxy_composite = 1) given a set of child and household context variables. This is a binary "
    "classification problem. The outcome is defined in Section 5.4.")

body(doc,
    "Critically, height-for-age z-score (HAZ) and weight-for-age z-score (WAZ) were explicitly excluded "
    "from the feature set, even though they are available as continuous measurements. The reason is that "
    "HAZ > −2 SD directly determines whether the physical proxy domain is passed — including HAZ as a "
    "predictor would create a circular situation where the model essentially predicts the outcome using "
    "the outcome itself. The honest question to ask is: which socioeconomic and demographic context "
    "variables can predict whether a child is nutritionally healthy AND regularly attending ECCE?")

heading2(doc, "7.2 Feature Set")

add_table(doc,
    ["Feature", "Description", "Type"],
    [
        ["child_age_months", "Age in months (continuous)", "Continuous"],
        ["is_female", "Sex: 1=female, 0=male", "Binary"],
        ["is_urban", "Residence: 1=urban, 0=rural", "Binary"],
        ["wealth_index", "Wealth quintile (1=poorest, 5=richest)", "Ordinal 1–5"],
        ["mother_edu", "Mother's education (0=none, 5=higher)", "Ordinal 0–5"],
        ["mother_higher_ed", "Mother has complete secondary or higher (1/0)", "Binary"],
        ["mother_age", "Mother's age in years", "Continuous"],
        ["st_* (35 dummies)", "State fixed effects (reference: Andaman & Nicobar Is.)", "Binary (one-hot)"],
    ]
)

body(doc,
    "Total features: 42 (7 individual-level + 35 state dummies). Missing values are imputed with "
    "feature-wise medians inside each cross-validation fold to prevent data leakage.")

heading2(doc, "7.3 Class Imbalance")

body(doc,
    "The outcome is imbalanced: 25.0% of children are 'on track' (Class 1 = 21,717) and 75.0% are "
    "'not on track' (Class 0 = 65,282). Without correction, a classifier that always predicts 'not on "
    "track' would achieve 75% accuracy but zero recall for Class 1. Both models use "
    "class_weight='balanced', which scales the loss function by the inverse class frequency:")
formula_para(doc, "w_class = n_samples / (n_classes × n_samples_in_class)")
body(doc,
    "This causes the model to penalise misclassifications of the minority class more heavily.")

heading2(doc, "7.4 Models")

heading3(doc, "7.4.1 Logistic Regression (Baseline)")

body(doc,
    "Logistic Regression models the log-odds of the outcome as a linear combination of features:")
formula_para(doc, "log[P(Y=1) / P(Y=0)] = β₀ + β₁x₁ + β₂x₂ + ... + βₚxₚ")
formula_para(doc, "P(Y=1 | X) = σ(Xβ) = 1 / (1 + exp(−Xβ))")

body(doc,
    "where σ(·) is the logistic sigmoid function. L2 regularisation is applied (C=1.0, meaning the "
    "regularisation penalty is λ = 1/C = 1.0). Features are standardised to zero mean and unit variance "
    "before fitting (StandardScaler inside the pipeline). Maximum iterations = 1,000.")

body(doc,
    "Logistic Regression serves as an interpretable baseline. Its coefficients directly indicate the "
    "direction (positive/negative) and relative magnitude of each feature's association with the outcome, "
    "holding other features constant.")

heading3(doc, "7.4.2 HistGradientBoostingClassifier (Main Model)")

body(doc,
    "HistGradientBoosting is scikit-learn's implementation of histogram-based gradient boosted decision "
    "trees, similar in spirit to XGBoost and LightGBM. It builds an ensemble of decision trees sequentially, "
    "where each tree corrects the residuals of the previous ensemble.")

body(doc, "The ensemble prediction is:")
formula_para(doc, "F_M(x) = Σₘ₌₁ᴹ  γₘ · hₘ(x)")

body(doc,
    "where hₘ(x) is the m-th decision tree and γₘ is the step size (learning_rate). Each tree is fitted "
    "to the negative gradient of the loss function (binary cross-entropy for classification):")
formula_para(doc, "L(y, F) = − [ y·log(σ(F)) + (1−y)·log(1−σ(F)) ]")

body(doc, "Hyperparameters used:")
add_table(doc,
    ["Parameter", "Value", "Role"],
    [
        ["max_iter", "300", "Number of trees in the ensemble"],
        ["max_depth", "5", "Maximum depth per tree (controls overfitting)"],
        ["learning_rate", "0.05", "Shrinkage: smaller = more conservative, requires more trees"],
        ["class_weight", "balanced", "Handles class imbalance (see Section 7.3)"],
        ["random_state", "42", "Reproducibility seed"],
    ]
)

body(doc,
    "HistGradientBoosting natively handles missing values — it learns the optimal split direction "
    "for NaN values at training time — so it does not require imputation. However, inside the CV "
    "pipeline, an imputer is included to standardise the data handling.")

heading2(doc, "7.5 Cross-Validation Strategy")

body(doc,
    "5-fold Stratified Cross-Validation (StratifiedKFold, sklearn) is used to evaluate model "
    "generalisation. Stratified folds preserve the class ratio in each fold, which is important "
    "for imbalanced datasets.")

body(doc, "The procedure:")
body(doc, "1. The data (n = 86,999) is partitioned into 5 equal-sized folds, each containing ~17,400 observations.")
body(doc, "2. In each iteration, 4 folds are used for training and 1 fold for testing.")
body(doc, "3. The pipeline (impute → scale → model) is fit exclusively on training data in each fold.")
body(doc, "4. Predictions are generated on the held-out test fold.")
body(doc, "5. Performance metrics (AUC, F1) are averaged across 5 folds.")

body(doc,
    "Out-of-Fold (OOF) predictions: for each child, the prediction was made by the model that "
    "was trained without that child's data fold. Concatenating all 5 sets of held-out predictions "
    "gives a full set of n=86,999 OOF probabilities that serve as honest estimates of generalisation "
    "performance (no data leakage).")

heading2(doc, "7.6 Evaluation Metrics")

heading3(doc, "7.6.1 ROC-AUC (Area Under the Receiver Operating Characteristic Curve)")

body(doc,
    "The ROC curve plots the True Positive Rate (TPR, Recall) against the False Positive Rate (FPR) "
    "at all classification thresholds. AUC measures the probability that the model ranks a randomly "
    "chosen positive example higher than a randomly chosen negative example:")
formula_para(doc, "AUC = P(score(positive) > score(negative))")
formula_para(doc, "TPR (Recall) = TP / (TP + FN)")
formula_para(doc, "FPR          = FP / (FP + TN)")

body(doc,
    "AUC = 0.5 corresponds to random guessing; AUC = 1.0 is perfect discrimination. "
    "AUC is threshold-independent and robust to class imbalance.")

heading3(doc, "7.6.2 F1 Score")

body(doc,
    "F1 is the harmonic mean of precision and recall for the positive (on-track) class:")
formula_para(doc, "Precision = TP / (TP + FP)")
formula_para(doc, "Recall    = TP / (TP + FN)")
formula_para(doc, "F1        = 2 × Precision × Recall / (Precision + Recall)")

body(doc,
    "F1 is sensitive to the classification threshold (0.5 is used here) and rewards balanced "
    "performance on both precision and recall. It is more informative than accuracy under "
    "class imbalance.")

heading2(doc, "7.7 Results")

add_table(doc,
    ["Model", "ROC-AUC (mean ± SD)", "F1 (mean ± SD)"],
    [
        ["Logistic Regression", "0.656 ± 0.003", "0.438 ± 0.003"],
        ["HistGradientBoosting", "0.656 ± 0.004", "0.434 ± 0.004"],
    ]
)

body(doc,
    "Both models achieve virtually identical AUC of 0.656, suggesting that the relationship between "
    "the context predictors and the proxy composite is approximately linear — gradient boosting's "
    "capacity to capture nonlinear interactions does not improve performance here. The AUC of 0.656 "
    "represents meaningful predictive power from socioeconomic context variables alone (AUC = 0.5 "
    "is chance; AUC > 0.7 is generally considered 'acceptable').")

body(doc,
    "Confusion matrix (OOF, HGB, threshold = 0.5): The model correctly identifies 60% of 'on track' "
    "children (recall = 0.60) while maintaining a precision of 0.34. The high recall for Class 1 "
    "relative to precision reflects the use of class_weight='balanced', which biases the model toward "
    "identifying positives even at the cost of more false positives.")

add_figure(doc, f"{FIG}/08_roc_curve.png",
    "Figure 7. Cross-validated ROC Curves — LR and HGB Models (out-of-fold predictions).")
add_figure(doc, f"{FIG}/08b_confusion_matrix.png",
    "Figure 8. Confusion Matrix — HistGradientBoosting (OOF, threshold = 0.5).")

heading2(doc, "7.8 SHAP Interpretability")

body(doc,
    "SHAP (SHapley Additive exPlanations) decomposes each model prediction into additive contributions "
    "from individual features, rooted in cooperative game theory. The SHAP value for feature i and "
    "observation x is defined as:")
formula_para(doc, "φᵢ(f, x) = Σ_{S ⊆ F\\{i}} [|S|!(|F|−|S|−1)! / |F|!] × [f(S∪{i}) − f(S)]")

body(doc,
    "where F is the full feature set and the sum is over all subsets S not containing feature i. "
    "Intuitively, φᵢ measures the average marginal contribution of feature i across all possible "
    "orderings of features. SHAP values satisfy axioms of efficiency (values sum to the model output "
    "minus the baseline), symmetry, linearity, and dummy (zero SHAP for irrelevant features).")

body(doc,
    "Implementation: shap.PermutationExplainer is used on a subsample of n=500 children "
    "(random seed 42), with max_samples=50 in the masker for computational efficiency. "
    "PermutationExplainer estimates SHAP values through random permutations of feature "
    "orderings rather than exact enumeration of all 2^p subsets.")

body(doc, "Top 5 SHAP predictors (mean |SHAP| across 500-row subsample):")

add_table(doc,
    ["Rank", "Feature", "Mean |SHAP|", "Interpretation"],
    [
        ["1", "Child age (months)", "0.031", "Older children in the 24–59 month range more likely on track"],
        ["2", "State: Meghalaya", "0.028", "Being in Meghalaya strongly reduces on-track probability (low ECCE access)"],
        ["3", "State: Nagaland", "0.025", "Being in Nagaland nearly eliminates on-track probability"],
        ["4", "Mother's education (ordinal)", "0.021", "Higher maternal education increases on-track probability"],
        ["5", "State: Uttar Pradesh", "0.020", "Being in UP reduces on-track probability (large, high-burden state)"],
    ]
)

add_figure(doc, f"{FIG}/09_shap_bar.png",
    "Figure 9. SHAP Feature Importance Bar Chart — Mean |SHAP Value| for each predictor.")
add_figure(doc, f"{FIG}/10_shap_beeswarm.png",
    "Figure 10. SHAP Beeswarm Plot — Direction and magnitude of each predictor's effect. "
    "Red = high feature value; Blue = low feature value. Positive SHAP = increases on-track probability.")

body(doc,
    "How to read the beeswarm (Figure 10): Each dot represents one child in the 500-row subsample. "
    "The horizontal position shows the SHAP value (effect on the model's output probability). "
    "Dots to the right (positive SHAP) push the prediction toward 'on track'; dots to the left push "
    "toward 'not on track'. Dot colour shows the actual feature value (red = high, blue = low). "
    "For example, for 'Mother's education': red dots (high education) appear on the right, indicating "
    "that higher maternal education increases the probability of being on track.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8: Clustering
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "8. Unsupervised Machine Learning: Clustering Developmental Profiles")

heading2(doc, "8.1 Motivation and Approach")

body(doc,
    "While supervised ML predicts a single outcome, unsupervised clustering asks: are there "
    "naturally occurring subgroups of children with similar developmental profiles? "
    "K-Means clustering is applied to group children based on their anthropometric z-scores "
    "(HAZ, WAZ) and ECCE attendance proxy (ecdi_proxy_learning), revealing distinct "
    "developmental profiles without imposing a predefined classification.")

heading2(doc, "8.2 K-Means Algorithm")

body(doc,
    "K-Means partitions n observations into k clusters by minimising the within-cluster sum of "
    "squared Euclidean distances (inertia/WCSS):")
formula_para(doc, "J(C) = Σₖ Σᵢ∈Cₖ ‖xᵢ − μₖ‖²")

body(doc,
    "where Cₖ is the set of points in cluster k and μₖ is the cluster centroid. The algorithm "
    "alternates between two steps until convergence:")
body(doc, "  (1) Assignment step: assign each point to the nearest centroid")
body(doc, "  (2) Update step: recompute each centroid as the mean of its assigned points")

body(doc,
    "Features are standardised (zero mean, unit variance) before clustering so that HAZ "
    "(range approximately −6 to +6) and the binary learning proxy (0 or 1) contribute equally.")

heading2(doc, "8.3 Choosing the Optimal k")

body(doc,
    "Two complementary criteria guide the choice of k:")

body(doc,
    "Elbow Method: Plot inertia (WCSS) against k. The 'elbow' where inertia reduction plateaus "
    "suggests a natural k. A subjective criterion — there is no universally agreed elbow detection "
    "algorithm.")

body(doc,
    "Silhouette Score: For each point, the silhouette coefficient measures how similar it is to "
    "its own cluster (cohesion) relative to the nearest other cluster (separation):")
formula_para(doc, "s(i) = [b(i) − a(i)] / max{a(i), b(i)}")

body(doc,
    "where a(i) = mean distance from point i to other points in the same cluster, and "
    "b(i) = mean distance from point i to points in the nearest different cluster. "
    "s(i) ranges from −1 (misclassified) to +1 (well-separated). The mean silhouette score "
    "across all points is maximised to select k.")

body(doc,
    "Computational note: Silhouette is O(n²) in time complexity. For n ≈ 87,000, this would "
    "take several hours. A subsample of 10,000 children (random seed 42) is therefore used "
    "for k-selection; the final model is fit on all 86,834 children with complete data.")

heading2(doc, "8.4 Results: k = 2 Optimal")

body(doc,
    "The silhouette score is maximised at k = 2. This result reflects the dominant role of the "
    "binary learning proxy: children either attend anganwadi regularly or they do not, and this "
    "binary split creates two well-separated groups. The anthropometric variables (HAZ, WAZ) show "
    "only small between-cluster differences.")

add_table(doc,
    ["Cluster", "Label", "n", "Physical On-Track", "Nutritional On-Track", "Learning On-Track", "Mean HAZ", "Mean WAZ"],
    [
        ["0", "Profile 1 (no regular ECCE)", "51,470", "61.5%", "67.5%", "0.0%", "−1.49", "−1.47"],
        ["1", "High ECCE Attendance", "35,364", "61.3%", "64.5%", "100.0%", "−1.53", "−1.57"],
    ]
)

body(doc,
    "Key finding: The two clusters have nearly identical anthropometric profiles (HAZ ≈ −1.5, "
    "WAZ ≈ −1.5 in both) but are perfectly separated by ECCE attendance. This demonstrates that "
    "in India, nutritional status and ECCE access are largely independent — a child's height does "
    "not predict whether they attend anganwadi, and attending anganwadi does not improve height "
    "(at least not within the cross-sectional timeframe of this survey). This has important policy "
    "implications: stunting and ECCE non-attendance are co-existing but distinct problems that "
    "require separate programmatic responses.")

heading2(doc, "8.5 Principal Component Analysis (PCA) Visualisation")

body(doc,
    "PCA reduces the 3-dimensional standardised feature space (HAZ, WAZ, ecdi_proxy_learning) "
    "to 2 dimensions for visualisation, preserving maximum variance. The transformation is:")
formula_para(doc, "Z = XW   where W ∈ ℝ^(3×2) contains the top 2 eigenvectors of the covariance matrix")

body(doc,
    "The percentage of variance explained by each principal component is shown on the axis labels "
    "of Figure 12. Because the learning proxy is binary and orthogonal to HAZ/WAZ, the PCA scatter "
    "shows two vertically/horizontally separated bands corresponding to the two clusters.")

add_figure(doc, f"{FIG}/11_cluster_selection.png",
    "Figure 11. Elbow Method (left) and Silhouette Score (right) for k = 2 to 8. "
    "Vertical red line marks the silhouette-optimal k.")
add_figure(doc, f"{FIG}/12_cluster_profiles.png",
    "Figure 12. Cluster Profiles — Proxy Domain On-Track Rates by Cluster.")
add_figure(doc, f"{FIG}/13_cluster_pca.png",
    "Figure 13. PCA Visualisation of Clusters (2,000-row subsample per cluster for speed). "
    "Black X marks = cluster centroids in PCA space.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9: Geospatial Analysis
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "9. Geospatial Analysis: State-Level Patterns and Spatial Autocorrelation")

heading2(doc, "9.1 State-Level Aggregation")

body(doc,
    "The proxy composite, physical proxy, and learning proxy on-track rates are computed for each "
    "state by aggregating individual-level outcomes. Only states with at least 50 children in the "
    "composite sample are included (all 36 states/UTs meet this criterion).")

heading2(doc, "9.2 Ranked Bar Charts")

body(doc,
    "Figure 14 (06_state_rates.png) presents a horizontal bar chart of all 36 states sorted by "
    "composite on-track rate. States above the national median are shown in blue; states below in "
    "red. The median is overlaid as a dashed vertical line.")

body(doc,
    "Figure 15 (06b_state_physical.png) presents the same chart for the physical proxy (non-stunting "
    "rate) alone. Comparing Figures 14 and 15 reveals interesting divergences: some states "
    "(e.g., Kerala) have very high non-stunting rates but low composite rates — reflecting low "
    "anganwadi attendance, likely due to private preschool substitution.")

heading2(doc, "9.3 Physical vs Learning Scatter Plot")

body(doc,
    "Figure 16 (07_state_scatter.png) plots each state as a point in the space of (physical proxy "
    "rate, learning proxy rate). Bubble size is proportional to sample size. States in the "
    "upper-right quadrant excel on both dimensions; states in the lower-left have challenges on both. "
    "Kerala exemplifies the 'high physical, low learning proxy' quadrant — high nutrition but low "
    "formal anganwadi attendance. West Bengal and Odisha appear in the upper-right, reflecting "
    "strong ICDS delivery alongside relatively lower stunting.")

heading2(doc, "9.4 Moran's I — Spatial Autocorrelation")

body(doc,
    "Moran's I measures the degree of spatial autocorrelation in the ECD rates — i.e., whether "
    "geographically neighbouring states tend to have similar rates. It is the spatial analogue of "
    "the Pearson correlation coefficient.")

body(doc, "The standard formulation using row-standardised spatial weights is:")
formula_para(doc, "I = (Σᵢ Σⱼ w*ᵢⱼ · zᵢ · zⱼ) / (Σᵢ zᵢ²)")

body(doc,
    "where zᵢ = yᵢ − ȳ is the mean-centred ECD rate for state i, and w*ᵢⱼ is the row-standardised "
    "spatial weight between states i and j:")
formula_para(doc, "w*ᵢⱼ = wᵢⱼ / Σⱼ wᵢⱼ   (each row of W sums to 1)")

body(doc,
    "For row-standardised weights, the normalisation constant n/S₀ = n/n = 1 (since the sum of all "
    "weights S₀ = n when each row sums to 1), so the formula simplifies as written above.")

body(doc,
    "The spatial weight matrix W is defined by land-border adjacency (wᵢⱼ = 1 if states i and j "
    "share a land border; 0 otherwise). This is hand-coded from geographic knowledge of India's "
    "state boundaries, covering 31 contiguous states. Island territories and small UTs without "
    "land borders are excluded from the Moran's I calculation.")

body(doc, "Statistical significance is assessed via a permutation test:")
body(doc, "1. The observed Moran's I is computed (I_obs).")
body(doc, "2. The ECD rates are randomly shuffled 999 times across states.")
body(doc, "3. For each permutation, Moran's I is recomputed (I_perm).")
body(doc, "4. The p-value is: p = (#{I_perm ≥ I_obs} + 1) / (999 + 1)")

body(doc, "Results:")
add_table(doc,
    ["Statistic", "Value"],
    [
        ["States in adjacency matrix", "31"],
        ["Moran's I", "0.421"],
        ["Permutation p-value", "0.001 (999 permutations)"],
        ["Interpretation", "Significant positive spatial clustering"],
    ]
)

body(doc,
    "Interpretation: Moran's I = 0.421 (p = 0.001) indicates strong, statistically significant "
    "positive spatial autocorrelation. Neighbouring states tend to have similar ECD proxy rates. "
    "This is consistent with regional clusters: high-performing southern and eastern states "
    "(West Bengal, Odisha, Andhra Pradesh, Gujarat) cluster together, while low-performing "
    "northern states (Uttar Pradesh, Bihar, Rajasthan) also cluster together. The Northeast "
    "cluster (Assam, Meghalaya, Nagaland, Manipur) shows uniformly low rates due to the "
    "anganwadi measurement issue described earlier.")

body(doc,
    "Figure 17 (07b_moran_scatter.png) shows the Moran scatter plot: each point is a state, "
    "plotted with its ECD rate on the x-axis and its spatially lagged rate (weighted mean of "
    "neighbouring states' rates) on the y-axis. The positive slope of the regression line "
    "visually confirms positive spatial autocorrelation. States in the top-right quadrant "
    "(high ECD rate, high neighbour rate) form a spatial cluster of 'winners'; states in "
    "the bottom-left form a cluster of 'laggards'.")

add_figure(doc, f"{FIG}/06_state_rates.png",
    "Figure 14. ECD Proxy Composite Rates by State (ranked). Blue = above median; Red = below.",
    width_in=5.0)
add_figure(doc, f"{FIG}/06b_state_physical.png",
    "Figure 15. Physical Proxy (Non-Stunting) Rate by State (ranked).",
    width_in=5.0)
add_figure(doc, f"{FIG}/07_state_scatter.png",
    "Figure 16. Physical vs Learning Proxy Rate Scatter — Each state as a bubble (size = sample size).")
add_figure(doc, f"{FIG}/07b_moran_scatter.png",
    "Figure 17. Moran Scatter Plot — ECD Proxy Rate vs Spatially Lagged Rate (row-standardised W). "
    "Positive slope confirms spatial autocorrelation (I = 0.421, p = 0.001).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10: Econometric Analysis
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "10. Econometric Analysis: Survey-Weighted, Clustered, Multilevel, "
              "Selection, and Decomposition")

body(doc,
    "Section 7's supervised ML pipeline produces predictions; this section conducts inferential "
    "econometric analysis on the expanded 4-domain ECD proxy composite (Section 5.6). Six specifications "
    "are reported, each addressing a distinct concern.")

heading2(doc, "10.1 Survey-Weighted Descriptives")
body(doc,
    "Sampling weights v005/1,000,000 from the NFHS-5 women's recode are applied to compute population-"
    "representative on-track rates. With the individual-score 4-domain composite, the weighted overall "
    "rate is 36.86% — close to the unweighted 35.82%.")

heading2(doc, "10.2 Cluster-Robust Logistic Regression (PSU Clustering)")
body(doc,
    "We estimate the binary outcome ecdi_proxy4_composite ∈ {0, 1} via a Generalised Linear Model with "
    "binomial family and logit link, including 35 state fixed-effect dummies (reference state: Andaman & "
    "Nicobar). Standard errors are clustered at the Primary Sampling Unit level (v001 — the cluster ID "
    "from the two-stage NFHS-5 sample design), correctly accounting for within-cluster correlation. "
    "The model uses 25,706 PSUs in the analytic sample of 86,999 children.")

add_table(doc,
    ["Predictor", "Coefficient", "Cluster SE", "z", "p"],
    [
        ["child_age_months",  "+0.014",  "0.001", "14.35", "< 0.001"],
        ["is_female",         "+0.047",  "0.020", " 2.31", "  0.021"],
        ["is_urban",          "+0.027",  "0.034", " 0.78", "  0.435"],
        ["wealth_index",      "+0.580",  "0.011", "50.60", "< 0.001"],
        ["mother_edu",        "+0.424",  "0.009", "47.71", "< 0.001"],
        ["mother_age",        "+0.037",  "0.002", "16.16", "< 0.001"],
    ]
)
body(doc,
    "McFadden Pseudo R² = 0.447 (with state dummies). Wealth, mother's education, and mother's age are "
    "the strongest individual-level predictors. A one-quintile increase in wealth raises the log-odds of "
    "being on track by 0.580 (≈ 79% increase in odds); one step up in maternal education raises log-odds "
    "by 0.424 (≈ 53% increase in odds). The is_urban coefficient is small (+0.027) and not significant in "
    "this PSU-clustered specification — once state, individual SES, and the new individual-level Literacy-"
    "Numeracy and Socio-emotional scores are controlled, raw urban residence does not add explanatory "
    "power. The earlier 'urban paradox' result (negative is_urban) under the binary state-merged proxies "
    "was partly an artefact of the imputation collapsing within-state heterogeneity; with individual-level "
    "scores the urban indicator is approximately a null effect.")

heading2(doc, "10.3 Average Marginal Effects (Probit)")
body(doc,
    "Logit coefficients are not directly interpretable as probability changes. We re-estimate the same "
    "specification as a probit and compute Average Marginal Effects (AMEs) — the expected change in the "
    "probability of being on track for a one-unit change in the predictor, averaged across the empirical "
    "distribution.")

add_table(doc,
    ["Predictor", "AME (Δ probability)", "Std. Err.", "p"],
    [
        ["child_age_months",  "+0.0016", "0.0001", "< 0.001"],
        ["is_female",         "+0.0059", "0.0023", "  0.010"],
        ["is_urban",          "+0.0016", "0.0032", "  0.618"],
        ["wealth_index",      "+0.0660", "0.0011", "< 0.001"],
        ["mother_edu",        "+0.0478", "0.0009", "< 0.001"],
        ["mother_age",        "+0.0040", "0.0002", "< 0.001"],
    ]
)
body(doc,
    "Substantively: each step up the wealth quintile raises the probability of being on track by 6.6 pp; "
    "each level of mother's education adds 4.8 pp; each year of mother's age adds 0.4 pp; girls are 0.6 pp "
    "more likely than boys; older children (within 24–59 months) are 0.16 pp more likely per month. The "
    "urban AME is +0.16 pp and not statistically distinguishable from zero (p = 0.62).")

heading2(doc, "10.4 Multilevel LPM — Children Nested in States")
body(doc,
    "To partition variance between within-state and between-state components, we estimate a two-level "
    "linear probability model with a state-level random intercept (sm.MixedLM with REML). Fixed effects "
    "include the same individual-level covariates plus the externally-merged state-level proxies "
    "(literacy_numeracy_state_rate from ASER 2019; socioemo_state_index from NFHS-5 factsheet). The "
    "random-intercept variance captures residual state heterogeneity NOT explained by the merged proxies.")

add_table(doc,
    ["Component", "Variance", "ICC contribution"],
    [
        ["Between-state (random intercept)", "0.01426", "ICC = 0.105"],
        ["Within-state (residual)",          "0.12140", "—"],
    ]
)
body(doc,
    "ICC = 0.105 means that ~10% of the residual variation in the ECD proxy composite (after controlling "
    "for individual covariates and the merged state-level proxies) lies BETWEEN states. Both state-level "
    "proxies are highly significant in this specification: literacy_numeracy_state_rate β = 1.925 "
    "(p < 0.001) and socioemo_state_index β = 0.641 (p < 0.001) — a 10-percentage-point increase in a "
    "state's ASER literacy rate raises a child's on-track probability by ~19 pp; a one-standard-deviation "
    "increase in the state's women's-empowerment composite raises it by ~6 pp. Most of the residual "
    "variance is within-state (90%), meaning individual-level variation now drives most of the unexplained "
    "outcome dispersion — consistent with the construction of individual-level Literacy-Numeracy and "
    "Socio-emotional scores in Section 5.")

heading2(doc, "10.5 Heckman Two-Step Selection Model")
body(doc,
    "Anganwadi attendance (s562) is observed only for children whose families enrolled in the ICDS programme "
    "(s558 = 'yes'). If enrollment is non-random with respect to the ECD outcome, regressions on the enrolled "
    "subsample suffer from selection bias. We address this with a Heckman two-step procedure:")
body(doc,
    "Selection equation: a probit predicts Pr(s558 = 'yes') from individual covariates plus the state-level "
    "anganwadi share (the exclusion restriction — varies across states but enters only the selection equation). "
    "Outcome equation: a linear probability model for the 4-domain composite among the enrolled, including the "
    "Inverse Mills Ratio (IMR) computed from the selection probit.")
body(doc,
    "The IMR coefficient in the outcome equation is −0.6249 (HC1 SE ≈ 0.017, p < 0.001), strongly "
    "significant and negative. Mechanically: unobserved characteristics that increase a household's "
    "propensity to enrol in anganwadi are NEGATIVELY associated with being ECD on-track. This is "
    "consistent with the policy targeting of ICDS — the programme is designed to enrol the most "
    "disadvantaged children, who are also most at risk of being below the on-track threshold. The naïve "
    "regression that ignores selection therefore underestimates the true effect of anganwadi attendance "
    "on ECD outcomes.")

heading2(doc, "10.6 Oaxaca-Blinder Decomposition — Urban-Rural Gap")
body(doc,
    "The unconditional urban-rural gap on the 4-domain composite is +22.6 pp. Decomposed (Neumark 1988):")

add_table(doc,
    ["Component", "pp", "% of total"],
    [
        ["Total gap (urban − rural)",                     "+22.56", "100.0%"],
        ["Endowments (different observable means)",       "+21.99", " 97.5%"],
        ["Unexplained / coefficient differences",          "+0.57",  "  2.5%"],
    ]
)
body(doc,
    "Interpretation: 97.5% of the urban-rural ECD gap is accounted for by differences in observable "
    "endowments — chiefly wealth, the merged state-level ASER literacy and socio-emotional proxies, "
    "and mother's education. The remaining 2.5% is the 'coefficient' component, which is essentially "
    "zero. Substantively: equalising urban and rural endowments would close almost the entire gap; "
    "the urban and rural ECD-production functions look essentially the same at fixed inputs.")

add_figure(doc, f"{FIG}/16_probit_AME.png",
    "Figure 16. Average Marginal Effects on P(on track) from the probit specification, "
    "with 95% confidence intervals (HC1 SEs).")
add_figure(doc, f"{FIG}/17_oaxaca.png",
    "Figure 17. Oaxaca-Blinder decomposition of the urban-rural ECD on-track gap (Neumark 1988 weighting).")

# ─────────────────────────────────────────────────────────────────────────────
# ROBUSTNESS SUITE (G–L)
# ─────────────────────────────────────────────────────────────────────────────

heading2(doc, "10.7 Robustness — Two-Way Clustering at PSU and State")
body(doc,
    "Sections 10.2–10.3 cluster standard errors at the PSU level. PSUs are nested within states, but "
    "shocks may also operate at the state level (e.g., correlated state-policy effects on ICDS delivery). "
    "We re-compute SEs using two-way clustering at PSU and state. As the PSU is fully nested within the "
    "state, the Cameron-Gelbach-Miller (2011) formula collapses to the larger of the two SEs; we report "
    "this conservative upper bound.")

add_table(doc,
    ["Predictor", "Coef", "SE (PSU)", "SE (state)", "SE (max)", "z (max)", "p (max)"],
    [
        ["child_age_months",  "+0.014", "0.001", "0.002", "0.002", " 6.12", "< 0.001"],
        ["is_female",         "+0.047", "0.020", "0.023", "0.023", " 2.06", "  0.040"],
        ["is_urban",          "+0.027", "0.034", "0.060", "0.060", " 0.45", "  0.653"],
        ["wealth_index",      "+0.580", "0.011", "0.075", "0.075", " 7.73", "< 0.001"],
        ["mother_edu",        "+0.424", "0.009", "0.055", "0.055", " 7.77", "< 0.001"],
        ["mother_age",        "+0.037", "0.002", "0.007", "0.007", " 4.95", "< 0.001"],
    ]
)
body(doc,
    "Wealth, mother's education, mother's age, child's age, and sex remain highly significant under the "
    "more conservative state-level clustering. The is_urban coefficient further loses significance "
    "(p = 0.65 under state clustering vs p = 0.43 under PSU clustering), reinforcing that urban residence "
    "has no robust effect once individual-level proxies are constructed.")

heading2(doc, "10.8 Robustness — Threshold × Composite-Rule Sweep")
body(doc,
    "The individual-score threshold (z = 0 ↔ above the median) and the composite k-of-4 rule are both "
    "user choices. We re-run the cluster-robust logit across a 3 × 3 grid of these choices and report "
    "the on-track rate, plus the wealth and is_urban coefficients in each cell:")
add_table(doc,
    ["Threshold (z)", "k of 4", "On-track %", "wealth β", "wealth p", "urban β", "urban p"],
    [
        ["−0.25", "2", "69.6", "+0.605", "< 0.001", "+0.289", "< 0.001"],
        ["−0.25", "3", "42.2", "+0.573", "< 0.001", "+0.071", "  0.036"],
        ["−0.25", "4", "12.4", "+0.354", "< 0.001", "−0.118", "  0.002"],
        [" 0.00", "2", "64.3", "+0.487", "< 0.001", "+0.231", "< 0.001"],
        [" 0.00", "3", "35.8", "+0.580", "< 0.001", "+0.027", "  0.435"],
        [" 0.00", "4",  "9.7", "+0.430", "< 0.001", "−0.140", "  0.001"],
        [" 0.25", "2", "57.7", "+0.427", "< 0.001", "+0.200", "< 0.001"],
        [" 0.25", "3", "29.5", "+0.568", "< 0.001", "+0.080", "  0.025"],
        [" 0.25", "4",  "7.6", "+0.460", "< 0.001", "−0.121", "  0.007"],
    ]
)
body(doc,
    "The wealth coefficient is highly stable (range 0.354 – 0.605, all p < 0.001). The is_urban "
    "coefficient flips sign between k = 2 (positive) and k = 4 (negative): the most-disadvantaged children "
    "needed for the strict 4-of-4 rule are concentrated in urban slums, while children passing the lenient "
    "2-of-4 rule are over-represented in better-off urban areas. The headline (k = 3, threshold = 0) cell "
    "shows urban ≈ 0, p = 0.43.")
add_figure(doc, f"{FIG}/19_robustness_sweep.png",
    "Figure 19. On-track % heatmap across thresholds (rows) and composite rules (columns). "
    "Central cell (z = 0, k = 3) is the headline specification.")

heading2(doc, "10.9 Robustness — IV / 2SLS for Anganwadi Attendance")
body(doc,
    "Individual anganwadi attendance is plausibly endogenous to unobserved household preferences for child "
    "investment. We instrument it with the state-level anganwadi share — variation in state-level ICDS "
    "supply that is plausibly orthogonal to a given household's preferences after controlling for individual "
    "and state-level covariates.")
add_table(doc,
    ["Stage", "Specification", "Coef", "SE", "p", "First-stage F"],
    [
        ["1st",   "state_anganwadi_share → ecdi_proxy_learning", "+0.998", "0.010", "< 0.001", "10,773.9"],
        ["OLS",   "y = β·attendance + …  (un-instrumented)",     "+0.293", "0.002", "< 0.001", "—"],
        ["2SLS",  "y = β·attendance_hat + …",                     "+0.374", "0.010", "< 0.001", "10,773.9"],
    ]
)
body(doc,
    "The first-stage F = 10,774 is far above the Stock-Yogo weak-instrument threshold (F > 10), confirming "
    "the instrument is strong. The 2SLS coefficient (+0.374) is substantially larger than the OLS "
    "coefficient (+0.293), implying that OLS UNDERESTIMATES the effect of anganwadi attendance on the ECD "
    "outcome — consistent with the negative selection bias detected in the Heckman model (Section 10.5). "
    "Caveat: the exclusion restriction (state-level supply only operates through individual attendance) "
    "is plausible but not testable; results should be read as one estimate among several.")

heading2(doc, "10.10 Robustness — Quantile Regression on Continuous Score")
body(doc,
    "OLS / probit / logit estimate effects at the conditional mean. With a continuous 4-domain ECD score "
    "(range 0–4) we estimate quantile regressions at τ ∈ {0.10, 0.25, 0.50, 0.75, 0.90} to check whether "
    "predictors operate uniformly across the conditional distribution.")
add_table(doc,
    ["Quantile τ", "wealth_index", "mother_edu", "is_urban", "is_female"],
    [
        ["0.10", "+0.187", "+0.208", "+0.098", "+0.018"],
        ["0.25", "+0.146", "+0.138", "+0.083", "+0.018"],
        ["0.50", "+0.175", "+0.197", "+0.068", "+0.038"],
        ["0.75", "+0.143", "+0.154", "+0.085", "+0.018"],
        ["0.90", "+0.153", "+0.152", "+0.041", "+0.017"],
    ]
)
body(doc,
    "Wealth and maternal education effects are very stable across quantiles (range 0.14–0.21). The "
    "is_urban effect is positive and largest at lower quantiles (0.10 ↔ +0.098, 0.90 ↔ +0.041), suggesting "
    "urban areas help the most for children at the bottom of the conditional ECD distribution. Quantile "
    "estimates reinforce the headline mean-effect findings.")
add_figure(doc, f"{FIG}/20_quantile_regression.png",
    "Figure 20. Quantile regression coefficients on the continuous 4-domain score, with 95% CIs. "
    "Wealth and education coefficients are stable across quantiles; urban has its largest effect at the "
    "bottom of the conditional distribution.")

heading2(doc, "10.11 Robustness — Drop the 12 Region-Imputed States")
body(doc,
    "The Literacy-Numeracy state-level ASER value is regionally imputed for 12 of 36 states/UTs. We re-fit "
    "the cluster-robust logit on the restricted sample of 80,475 children in the 24 states with direct "
    "ASER 2019 sampling. Coefficients are essentially unchanged:")
add_table(doc,
    ["Predictor", "Full sample β", "ASER-only β", "Difference"],
    [
        ["wealth_index",     "+0.580", "+0.567", "−0.013"],
        ["mother_edu",       "+0.424", "+0.422", "−0.002"],
        ["mother_age",       "+0.037", "+0.037", " 0.000"],
        ["is_urban",         "+0.027", "+0.048", "+0.021"],
        ["is_female",        "+0.047", "+0.055", "+0.008"],
        ["child_age_months", "+0.014", "+0.016", "+0.002"],
    ]
)
body(doc,
    "The headline effects (wealth, education, age) move by less than 0.02 log-odds units — within rounding. "
    "Findings are robust to dropping the regionally-imputed states.")

heading2(doc, "10.12 Robustness — Wild-Cluster Bootstrap on is_urban (B = 999, state cluster)")
body(doc,
    "With only 36 state clusters, asymptotic state-clustered SEs may be unreliable. We compute a wild-"
    "cluster-bootstrap p-value on the is_urban coefficient (Cameron-Gelbach-Miller 2008): refit the LPM, "
    "multiply within-state residuals by Rademacher draws (±1) per state, refit B = 999 times, collect "
    "the t-statistic distribution under the null.")
add_table(doc,
    ["Statistic", "Value"],
    [
        ["β (is_urban)",          "+0.005"],
        ["HC0 SE",                "0.004"],
        ["t observed",            "1.405"],
        ["WCB p (B = 999, two-sided)", "0.683"],
        ["Empirical 95% CI for t under H0", "[−3.88, +6.91]"],
    ]
)
body(doc,
    "Under proper state-level inference, the urban effect is unambiguously NOT significant (p = 0.683). "
    "This further confirms 10.7's finding: the urban indicator has no robust effect once individual-level "
    "proxies are constructed and SEs are adjusted for the few-clusters problem.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 11: Figures Explained (Summary)
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "11. Summary of All Figures")

add_table(doc,
    ["Figure", "Filename", "What It Shows", "Key Takeaway"],
    [
        ["1", "01_proxy_domain_rates.png",         "2-domain proxy on-track rates (legacy)",       "Composite 25% under the 2-domain rule"],
        ["2", "02_anthropometric_distributions.png","HAZ + WAZ histograms",                         "Both distributions shifted left; mean ≈ −1.47 SD"],
        ["3", "03_composite_by_wealth.png",        "2-domain composite by wealth quintile",        "Wealth gradient up to richer quintile"],
        ["6", "06_state_rates.png",                "2-domain state ranking",                       "8× spread across states"],
        ["7", "06b_state_physical.png",            "Physical proxy (non-stunting) by state",       "Kerala best; Bihar / MP / Meghalaya worst"],
        ["8", "07_state_scatter.png",              "Physical vs learning proxy by state",          "No simple correlation; regional clustering"],
        ["9", "07b_moran_scatter.png",             "Moran scatter — spatial autocorrelation",      "I = 0.421, p = 0.001"],
        ["10","08_roc_curve.png",                  "5-fold CV ROC for LR + HGB (legacy 2-domain)", "AUC ≈ 0.66 from context variables"],
        ["11","08b_confusion_matrix.png",          "HGB confusion matrix",                         "Recall-favoured under class_weight=balanced"],
        ["12","09_shap_bar.png",                   "SHAP feature importance",                      "Child age, state, mother's education on top"],
        ["13","10_shap_beeswarm.png",              "SHAP beeswarm",                                 "Direction + magnitude of effects"],
        ["14","11_cluster_selection.png",          "K-Means elbow + silhouette",                    "k = 2 silhouette-optimal"],
        ["15","12_cluster_profiles.png",           "Cluster proxy-rate profiles",                   "Split is anganwadi vs no anganwadi"],
        ["16","13_cluster_pca.png",                "PCA of K-Means clusters",                       "2 bands — ECCE binary dominates"],
        ["17","14_expanded_domain_rates.png",      "Expanded 4-domain proxy on-track rates (individual scores)", "Composite ≥3/4 = 35.8%"],
        ["18","15_state_rates_4dom.png",           "4-domain composite by state",                   "TN / Kerala / AP top; Bihar / UP / Rajasthan bottom"],
        ["19","16_probit_AME.png",                 "Probit AME forest plot",                        "Wealth + mother education + age are largest"],
        ["20","17_oaxaca.png",                     "Oaxaca-Blinder urban-rural decomposition",      "97.5% explained by endowments"],
        ["21","18_within_state_litnum.png",        "Within-state distribution of individual lit-num score", "Non-zero variation — ecological-fallacy mitigated"],
        ["22","19_robustness_sweep.png",           "Threshold × composite-rule heatmap",            "On-track stable in central cell, varies at extremes"],
        ["23","20_quantile_regression.png",        "Quantile regression on continuous ECD score",   "Wealth + education effects flat across quantiles"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 12: Limitations
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "12. Limitations")

limitations = [
    ("No direct individual-level cognitive or socio-emotional assessment",
     "NFHS-5 contains no items equivalent to ECDI2030 EC6–EC15. The Literacy-Numeracy and Socio-emotional "
     "domains are constructed via individual-level scores that combine a state-level signal (ASER 2019 "
     "or NFHS-5 factsheet) with a weighted sum of individual NFHS-5 covariates — see Section 5.4–5.5. "
     "Within-state SD of these scores is 0.42 and 0.39 respectively (non-zero by construction), so the "
     "ecological-fallacy concern of the previous binary state-merged version is mitigated. The 4-domain "
     "composite remains a hybrid individual-+-state construct rather than a direct ECDI2030 measurement; "
     "the state-level signal carries unit weight in the score, individual covariates carry 0.10–0.40."),
    ("State-level signal still dominates two domains",
     "Even with the individual-level score construction, the state-level ASER and NFHS-5-factsheet signals "
     "remain the largest single component of the Literacy-Numeracy and Socio-emotional scores (β = 1 vs "
     "0.10–0.40 for individual covariates). The proxy will systematically miss within-state heterogeneity "
     "that is uncorrelated with the chosen individual covariates (e.g., quality of caregiver-child verbal "
     "interaction, household library access). A fully individual-level proxy would require direct ECDI2030 "
     "administration in NFHS-5."),
    ("ASER 2019 covers 24 of 36 states; 12 imputed by region",
     "ASER 2019 Early Years sampled only 26 districts in 24 states. The 12 unsampled states/UTs (Andaman "
     "& Nicobar, Arunachal Pradesh, Chandigarh, Dadra & Nagar Haveli & Daman & Diu, Goa, Jammu & Kashmir, "
     "Ladakh, Lakshadweep, Mizoram, NCT of Delhi, Puducherry, Sikkim) receive their region's ASER mean "
     "as an imputed value rather than a fabricated one, but this introduces measurement noise for those "
     "states. ASER values themselves come from one rural district per sampled state, so within-state "
     "urban areas are not represented in the proxy."),
    ("Anganwadi selection bias (now quantified)",
     "The learning proxy is only available for children enrolled in ICDS (s558 = 'yes'). The Heckman "
     "two-step model in Section 10.5 confirms this is non-random selection (IMR coefficient = −0.6249, "
     "p < 0.001) and the inferential models are corrected accordingly, but children entirely outside the "
     "ICDS system are still excluded from the analytic sample."),
    ("HAZ as proxy for Physical domain",
     "HAZ measures accumulated nutritional deficits (linear growth faltering) rather than fine motor "
     "skills (EC9) or energy levels (EC10). A child can have HAZ > −2 SD and still have motor "
     "developmental delays."),
    ("Cross-sectional design",
     "NFHS-5 is cross-sectional. The econometric models in Section 10 estimate associations under unit-"
     "homogeneity assumptions; they do not identify causal effects. Anganwadi attendance, in particular, "
     "is plausibly endogenous to unobserved household preferences for child investment."),
    ("Northeast states anomaly",
     "Nagaland, Manipur, and Meghalaya show very low anganwadi attendance in NFHS-5, possibly "
     "reflecting alternative ECD delivery channels (church crèches, state programmes) not captured by "
     "the s562 question. The state-level proxy merge partially compensates but the Learning-domain "
     "component of the composite remains under-counted in these states."),
    ("SHAP subsample",
     "SHAP values are computed on a random subsample of 500 children due to PermutationExplainer "
     "compute cost. Results may not fully represent the full distribution of predictor effects."),
]

for title_lim, text_lim in limitations:
    p = doc.add_paragraph()
    run_bold = p.add_run(f"{title_lim}: ")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_text = p.add_run(text_lim)
    run_text.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 13: Conclusions
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "13. Conclusions and Policy Implications")

heading2(doc, "13.1 Summary of Findings")

body(doc,
    "This analysis of 133,750 NFHS-5 children integrates external state-level survey data — ASER 2019 "
    "Early Years (parsed directly from the public PDF) and the NFHS-5 published State Factsheet "
    "Compendium (via jvargh7/nfhs5_factsheets) — into individual-level Literacy-Numeracy and Socio-"
    "emotional scores. Each score combines a state-level signal with a weighted sum of within-state "
    "individual covariates (mother's education, wealth, urban, mother's age), preserving non-zero "
    "within-state variation (SD = 0.42 and 0.39) and mitigating the ecological-fallacy critique. "
    "Approximately 35.8% of children aged 24–59 months are classified on track (≥ 3 of 4 domains, "
    "ECDI2030 official rule); the survey-weighted figure is 36.9%.")

body(doc,
    "The core econometric specifications confirm that wealth, maternal education, and mother's age are "
    "the dominant individual-level predictors. A one-quintile increase in wealth raises the probability "
    "of being on track by 6.6 pp (probit AME); one step up in maternal education by 4.8 pp; one extra "
    "year of mother's age by 0.4 pp. Multilevel ICC = 0.105 — most of the residual variation now lies "
    "WITHIN states (90%), reflecting the individual-level construction of the new domain scores. "
    "The Heckman selection model confirms (IMR = −0.62, p < 0.001) that ICDS programme reach is "
    "appropriately targeting more disadvantaged children, and the IV / 2SLS estimate of the anganwadi "
    "treatment effect (β = +0.374) exceeds the OLS estimate (β = +0.293), reinforcing the negative-"
    "selection finding. The Oaxaca-Blinder decomposition shows that 97.5% of the urban-rural ECD gap "
    "(+22.6 pp) is accounted for by endowment differences alone — equalising endowments would essentially "
    "close the gap.")
body(doc,
    "The robustness suite (Sections 10.7–10.12) shows that the headline wealth, education, and age "
    "effects are stable across two-way clustering, threshold and composite-rule sweeps, IV identification, "
    "quantile regressions, dropping the regionally-imputed states, and a wild-cluster bootstrap. The is_"
    "urban coefficient is the only result that is NOT robust: positive at lenient k = 2 thresholds, near-"
    "zero at the headline k = 3 specification, and negative at strict k = 4. Once SEs are properly adjusted "
    "for the few-state-clusters problem (wild-cluster bootstrap, B = 999), urban residence has no robust "
    "effect on ECD outcomes — the apparent 'urban premium' is fully explained by endowments.")

heading2(doc, "13.2 Policy Implications")

policy_points = [
    ("Target dual deprivation",
     "Children who are both stunted AND not attending anganwadi represent the highest-risk group. "
     "ICDS programmes should actively identify and enrol such children through community-level screening."),
    ("Address ECCE access in the North",
     "Uttar Pradesh, Bihar, and Rajasthan show both high stunting rates AND low regular anganwadi "
     "attendance. Strengthening AWC attendance incentives (e.g., conditional nutrition supplements) "
     "in these states could improve both dimensions simultaneously."),
    ("Maternal education as a multiplier",
     "Mother's education is among the top SHAP predictors. Investments in female secondary education "
     "are likely to yield ECD dividends in the next generation."),
    ("Improve ECCE measurement in the Northeast",
     "NFHS-5's anganwadi-centric learning proxy fails to capture the ECD programmes delivered through "
     "non-ICDS channels in Nagaland, Manipur, and Meghalaya. Future surveys should include broader "
     "ECCE attendance questions."),
    ("Spatial targeting",
     "The significant spatial clustering (I=0.421) supports geographically concentrated interventions. "
     "Resources directed at state clusters that lag behind their neighbours would be efficient from "
     "a spillover perspective."),
]
for title_pol, text_pol in policy_points:
    p = doc.add_paragraph()
    run_bold = p.add_run(f"{title_pol}: ")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_text = p.add_run(text_pol)
    run_text.font.size = Pt(11)

heading2(doc, "13.3 Future Work")

body(doc,
    "1. If a future round of India MICS6 (with direct ECDI2030 administration) becomes available, "
    "the proxy composite can be benchmarked against direct measurement in matched districts.\n"
    "2. Survival/panel models using NFHS-4 (2015–16) and NFHS-5 (2019–21) together could capture "
    "within-state trends in ECD outcomes.\n"
    "3. Incorporating district-level data (NFHS-5 provides district-level estimates for 707 districts) "
    "would allow higher-resolution geospatial analysis and more granular ASER merging.\n"
    "4. Three-level mixed-effects models (children → district → state) would more fully partition "
    "variance and accommodate within-district correlation, which the present two-level (state) "
    "specification cannot.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

heading1(doc, "References")

references = [
    "Black, R.E., Victora, C.G., Walker, S.P., et al. (2013). Maternal and child undernutrition and "
    "overweight in low-income and middle-income countries. The Lancet, 382(9890), 427–451.",

    "Heckman, J.J. (2006). Skill formation and the economics of investing in disadvantaged children. "
    "Science, 312(5782), 1900–1902.",

"International Institute for Population Sciences (IIPS) & ICF. (2021). National Family Health Survey "
    "(NFHS-5), India, 2019–21. Mumbai: IIPS.",

    "Britto, P.R., Lye, S.J., Proulx, K., et al. (2017). Nurturing care: promoting early childhood "
    "development. The Lancet, 389(10064), 91–102. https://doi.org/10.1016/S0140-6736(16)31390-3",

    "UNICEF. (2023). Early Childhood Development Index 2030 (ECDI2030): Technical Manual. "
    "UNICEF Data and Analytics, New York.",

    "Victora, C.G., Adair, L., Fall, C., et al. (2008). Maternal and child undernutrition: consequences "
    "for adult health and human capital. The Lancet, 371(9609), 340–357.",

    "WHO Multicentre Growth Reference Study Group. (2006). WHO Child Growth Standards: Length/height-for-age, "
    "weight-for-age, weight-for-length, weight-for-height and body mass index-for-age — Methods and "
    "development. WHO Press, Geneva. ISBN: 9789241546935",

    "Lundberg, S.M., & Lee, S.I. (2017). A unified approach to interpreting model predictions. "
    "Advances in Neural Information Processing Systems, 30.",

    "Anselin, L. (1995). Local Indicators of Spatial Association — LISA. "
    "Geographical Analysis, 27(2), 93–115.",
]

for ref in references:
    p = doc.add_paragraph(ref, style="List Bullet")
    p.runs[0].font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/prachi/india_nfhs5_project/India_NFHS5_ECD_Methodology_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
